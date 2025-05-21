import os
import openai
import pdfplumber
import pytesseract
from PIL import Image
from docx import Document
from odf.opendocument import load
from odf.text import P

openai.api_key = os.getenv("OPENAI_API_KEY")

FOLDERS = ["Cywilna", "Apelacja", "Egzekucja"]

def extract_text_from_pdf(path):
    text = ""
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text += page_text + "\n"
                else:
                    # OCR dla stron bez tekstu
                    img = page.to_image(resolution=300)
                    ocr_text = pytesseract.image_to_string(img.original)
                    if ocr_text.strip():
                        text += ocr_text + "\n"
    except Exception as e:
        text = f"(Błąd odczytu PDF: {e})"
    return text

def extract_text_from_docx(path):
    try:
        doc = Document(path)
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        return f"(Błąd odczytu DOCX: {e})"

def extract_text_from_doc(path):
    # Obsługa starszego formatu .doc (zakładamy, że jest konwertowany do tekstu)
    try:
        doc = Document(path)  # python-docx może czasem obsłużyć .doc, jeśli jest kompatybilny
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        return f"(Błąd odczytu DOC: {e})"

def extract_text_from_odt(path):
    try:
        odt = load(path)
        paragraphs = []
        for elem in odt.getElementsByType(P):
            paragraphs.append(str(elem))
        return "\n".join(paragraphs)
    except Exception as e:
        return f"(Błąd odczytu ODT: {e})"

def extract_text_from_txt(path):
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    except Exception as e:
        return f"(Błąd odczytu TXT: {e})"

def summarize(text, filename, folder):
    prompt = (
        f"Jesteś elitarnym zespołem prawników AI Cyfrowej Kancelarii Wiktora Andrukiewicza z Wrocławia. "
        f"Streszcz poniższy dokument ({filename}) z folderu {folder} w 5-10 zdaniach, podkreślając kluczowe fakty, argumenty i znaczenie dla sprawy. "
        f"Zachowaj styl profesjonalny i zgodny z wytycznymi kancelarii. Oto treść dokumentu:\n\n{text[:12000]}"
    )
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4-turbo",
            messages=[{"role": "system", "content": prompt}],
            max_tokens=600
        )
        return response.choices[0].message["content"].strip()
    except Exception as e:
        return f"(Błąd generowania streszczenia: {e})"

for folder in FOLDERS:
    folder_path = os.path.join(os.getcwd(), folder)
    if not os.path.isdir(folder_path):
        continue

    print(f"Przetwarzam folder: {folder}")
    summaries = []
    for fname in os.listdir(folder_path):
        if fname.lower() in ["knowledge.md", "reader.txt", "readme.md"]:
            continue
        fpath = os.path.join(folder_path, fname)
        if not os.path.isfile(fpath):
            continue

        print(f"Znaleziono plik: {fname}")
        # Sprawdź, czy streszczenie już istnieje w knowledge.md - tymczasowo wyłączone dla testów
        summary_tag = f"### {fname}"
        knowledge_path = os.path.join(folder_path, "knowledge.md")
        # if os.path.exists(knowledge_path):
        #     with open(knowledge_path, "r", encoding="utf-8") as kf:
        #         if summary_tag in kf.read():
        #             print(f"Streszczenie dla {fname} już istnieje, pomijam.")
        #             continue  # już podsumowane

        # Ekstrakcja tekstu
        if fname.lower().endswith(".pdf"):
            text = extract_text_from_pdf(fpath)
            print(f"Wygenerowano tekst z PDF: {fname}, długość: {len(text)} znaków")
        elif fname.lower().endswith(".docx"):
            text = extract_text_from_docx(fpath)
            print(f"Wygenerowano tekst z DOCX: {fname}, długość: {len(text)} znaków")
        elif fname.lower().endswith(".doc"):
            text = extract_text_from_doc(fpath)
            print(f"Wygenerowano tekst z DOC: {fname}, długość: {len(text)} znaków")
        elif fname.lower().endswith(".odt"):
            text = extract_text_from_odt(fpath)
            print(f"Wygenerowano tekst z ODT: {fname}, długość: {len(text)} znaków")
        elif fname.lower().endswith(".txt"):
            text = extract_text_from_txt(fpath)
            print(f"Wygenerowano tekst z TXT: {fname}, długość: {len(text)} znaków")
        else:
            print(f"Nieobsługiwany format pliku: {fname}, pomijam.")
            continue

        # Generowanie streszczenia
        summary = summarize(text, fname, folder)
        print(f"Wygenerowano streszczenie dla: {fname}")
        summaries.append(f"### {fname}\n{summary}\n")

    # Dopisz nowe streszczenia do knowledge.md
    if summaries:
        with open(os.path.join(folder_path, "knowledge.md"), "a", encoding="utf-8") as kf:
            kf.write("\n".join(summaries))
            print(f"Dopisano {len(summaries)} streszczeń do {knowledge_path}")
    else:
        print(f"Brak nowych streszczeń do dopisania w folderze {folder}")
